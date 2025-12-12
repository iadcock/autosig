#!/usr/bin/env python3
"""
Generate Word (.docx) reports from trading bot logs.

Usage:
    python report_docx.py              # Last 24 hours (default)
    python report_docx.py --hours 48   # Override hours
"""

import argparse
import os
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List, Any

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from jsonl_logger import get_raw_alerts, get_parsed_alerts, get_execution_plans

REPORTS_DIR = Path("reports")


def ensure_reports_dir():
    """Ensure reports directory exists."""
    REPORTS_DIR.mkdir(exist_ok=True)


def generate_report(hours: int = 24) -> str:
    """
    Generate a Word report for the last N hours.
    Returns the path to the generated report.
    """
    ensure_reports_dir()
    
    raw_alerts = get_raw_alerts(hours)
    parsed_alerts = get_parsed_alerts(hours)
    execution_plans = get_execution_plans(hours)
    
    parsed_by_id = {p["post_id"]: p for p in parsed_alerts}
    plans_by_id = {p["post_id"]: p for p in execution_plans}
    
    signals = [p for p in parsed_alerts if p.get("classification") == "SIGNAL"]
    non_signals = [p for p in parsed_alerts if p.get("classification") == "NON_SIGNAL"]
    
    executed = [p for p in execution_plans if p.get("action") == "PLACE_ORDER"]
    skipped = [p for p in execution_plans if p.get("action") == "SKIP"]
    closed = [p for p in execution_plans if p.get("action") == "CLOSE_ORDER"]
    
    doc = Document()
    
    title = doc.add_heading("Whop Trade Alerts Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    now = datetime.utcnow()
    start_time = now - timedelta(hours=hours)
    
    doc.add_paragraph(f"Report Generated: {now.strftime('%Y-%m-%d %H:%M:%S')} UTC")
    doc.add_paragraph(f"Time Range: {start_time.strftime('%Y-%m-%d %H:%M')} to {now.strftime('%Y-%m-%d %H:%M')} UTC ({hours} hours)")
    doc.add_paragraph()
    
    doc.add_heading("Summary Metrics", level=1)
    
    summary_table = doc.add_table(rows=6, cols=2)
    summary_table.style = 'Table Grid'
    
    metrics = [
        ("Total Raw Alerts", len(raw_alerts)),
        ("Parsed Signals", len(signals)),
        ("Non-Signal Alerts", len(non_signals)),
        ("Orders Placed", len(executed)),
        ("Orders Skipped", len(skipped)),
        ("Positions Closed", len(closed)),
    ]
    
    for i, (label, value) in enumerate(metrics):
        row = summary_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)
    
    doc.add_paragraph()
    
    if not raw_alerts and not parsed_alerts:
        doc.add_heading("No Alerts Found", level=1)
        doc.add_paragraph(f"No alerts were recorded in the last {hours} hours.")
    else:
        if signals:
            doc.add_heading("Trading Signals", level=1)
            doc.add_paragraph(f"Found {len(signals)} trading signal(s) in the reporting period.")
            doc.add_paragraph()
            
            signal_table = doc.add_table(rows=1, cols=6)
            signal_table.style = 'Table Grid'
            
            header_cells = signal_table.rows[0].cells
            headers = ["Ticker", "Strategy", "Expiry", "Legs", "Limit Range", "Status"]
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            for sig in signals:
                ps = sig.get("parsed_signal", {}) or {}
                plan = plans_by_id.get(sig.get("post_id"), {})
                
                row = signal_table.add_row()
                row.cells[0].text = ps.get("ticker", "N/A")
                row.cells[1].text = ps.get("strategy", "N/A")
                row.cells[2].text = ps.get("expiry", "N/A")
                
                legs = ps.get("legs", [])
                if legs:
                    legs_str = "; ".join([f"{l.get('side', '?')} {l.get('strike', '?')} {l.get('option_type', '?')}" for l in legs])
                else:
                    legs_str = "N/A"
                row.cells[3].text = legs_str[:50]
                
                limit_min = ps.get("limit_min", "?")
                limit_max = ps.get("limit_max", "?")
                row.cells[4].text = f"${limit_min} - ${limit_max}"
                
                action = plan.get("action", "UNKNOWN")
                row.cells[5].text = action
            
            doc.add_paragraph()
        
        if signals:
            doc.add_heading("Execution Plans", level=1)
            
            for sig in signals:
                ps = sig.get("parsed_signal", {}) or {}
                post_id = sig.get("post_id", "unknown")
                plan = plans_by_id.get(post_id, {})
                
                ticker = ps.get("ticker", "Unknown")
                strategy = ps.get("strategy", "Unknown")
                
                doc.add_heading(f"{ticker} - {strategy}", level=2)
                
                order_preview = plan.get("order_preview", {}) or {}
                
                doc.add_paragraph(f"Post ID: {post_id}")
                doc.add_paragraph(f"Action: {plan.get('action', 'N/A')}")
                doc.add_paragraph(f"Reason: {plan.get('reason', 'N/A')}")
                doc.add_paragraph(f"DRY_RUN: {plan.get('DRY_RUN', 'N/A')}")
                doc.add_paragraph(f"LIVE_TRADING: {plan.get('LIVE_TRADING', 'N/A')}")
                
                if order_preview:
                    doc.add_paragraph("Order Details:")
                    details = []
                    if "quantity" in order_preview:
                        details.append(f"  Contracts: {order_preview['quantity']}")
                    if "limit_price" in order_preview:
                        details.append(f"  Limit Price: ${order_preview['limit_price']}")
                    if "order_type" in order_preview:
                        details.append(f"  Order Type: {order_preview['order_type']}")
                    if "side" in order_preview:
                        details.append(f"  Side: {order_preview['side']}")
                    if "legs" in order_preview:
                        details.append(f"  Legs: {len(order_preview['legs'])}")
                    for d in details:
                        doc.add_paragraph(d)
                
                doc.add_paragraph()
        
        if non_signals:
            doc.add_heading("Non-Signal Alerts", level=1)
            doc.add_paragraph(f"Found {len(non_signals)} non-signal alert(s) that were not traded.")
            doc.add_paragraph()
            
            for ns in non_signals:
                reason = ns.get("non_signal_reason", "Not a trading signal")
                excerpt = ns.get("raw_excerpt", "")[:200]
                
                doc.add_heading(f"Non-Signal: {reason}", level=3)
                if excerpt:
                    p = doc.add_paragraph()
                    p.add_run("Excerpt: ").bold = True
                    p.add_run(excerpt + "..." if len(ns.get("raw_excerpt", "")) > 200 else excerpt)
                doc.add_paragraph()
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"whop_trade_report_{timestamp}.docx"
    filepath = REPORTS_DIR / filename
    
    doc.save(filepath)
    
    return str(filepath)


def main():
    parser = argparse.ArgumentParser(
        description="Generate Word report from trading bot logs"
    )
    parser.add_argument(
        "--hours",
        type=int,
        default=24,
        help="Number of hours to include in report (default: 24)"
    )
    
    args = parser.parse_args()
    
    print(f"Generating report for the last {args.hours} hours...")
    filepath = generate_report(args.hours)
    print(f"Report saved to: {filepath}")


if __name__ == "__main__":
    main()
